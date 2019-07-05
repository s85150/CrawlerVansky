from __future__ import print_function
import httplib2
import os
import io

import requests
from requests_html import HTML
import urllib.parse
import docx
from docx.dml.color import ColorFormat
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX

from apiclient import discovery
from oauth2client import client
from oauth2client import tools
from oauth2client.file import Storage
from apiclient.http import MediaFileUpload, MediaIoBaseDownload

import time
import datetime
import re
from multiprocessing import Pool
from bs4 import BeautifulSoup

try:
  import argparse
  flags = argparse.ArgumentParser(parents=[tools.argparser]).parse_args()
except ImportError:
  flags = None

SCOPES = 'https://www.googleapis.com/auth/drive/Canada_Rent_House'
CLIENT_SECRET_FILE = 'client_id.json'
APPLICATION_NAME = 'CrawlerVansky'


def get_credentials():
  """取得有效的憑證
     若沒有憑證，或是已儲存的憑證無效，就會自動取得新憑證

     傳回值：取得的憑證
  """
  credential_path = os.path.join("./", 'google-ocr-credential.json') #相對路徑
  store = Storage(credential_path)
  credentials = store.get()
  if not credentials or credentials.invalid:
    flow = client.flow_from_clientsecrets(CLIENT_SECRET_FILE, SCOPES)
    flow.user_agent = APPLICATION_NAME
    if flags:
      credentials = tools.run_flow(flow, store, flags)
    else: # Needed only for compatibility with Python 2.6
      credentials = tools.run(flow, store)
    print('憑證儲存於：' + credential_path)
  return credentials


def google_upload(file_name):

  # 取得憑證、認證、建立 Google 雲端硬碟 API 服務物件
  credentials = get_credentials()
  http = credentials.authorize(httplib2.Http())
  service = discovery.build('drive', 'v3', http=http)

  txtfile = file_name
  parents = '1vNSHZcNswxLxyUDbwjN4-wwXhqahYeBr'

  mime = 'application/vnd.google-apps.document'
  res = service.files().create(
    body={
      'name': txtfile,
      'parents': [parents],
      'mimeType': mime
    },
    media_body=MediaFileUpload(txtfile, mimetype=mime, resumable=True)
  ).execute()


def fetch(url):
    response = requests.get(url)
    return response


def parse_article_entries(doc):
    html = HTML(html=doc)
    soup = BeautifulSoup(html.text, "lxml")
    #post_entries = html.find('tr', class_='freeAdPadding')
    post_entries = html.find('tr.freeAdPadding')
    return post_entries

def parse_article_content_entries(doc):
    html = HTML(html=doc)
    post_entries = html.find('html')
    return post_entries


def parse_article_meta(entry):
    #print(entry)
    return {
        'title': entry.find('td a.adsTitleFont', first=True).text,
        'link': entry.find('td > a.adsTitleFont', first=True).attrs['href'],
        'img': entry.find('td > img', first=True).attrs['data-src'],
    }


def parse_article_content_meta(entry):
    return {
        'content': entry.find('div.cell.ctent', first=True).text,
        'content_replace_p': entry.find('div.cell.ctent p'),
        'time': entry.find('div.col-md-12.content-time', first=True).text,
    }


def get_metadata_from(url):
    resp = fetch(url)
    post_entries = parse_article_entries(resp.text)

    #metadata = [parse_article_meta(entry) for entry in post_entries]
    metadata = []
    for entry in post_entries:
      print(entry.attrs['itemtype'])
      if entry.attrs['itemtype'] == 'http://schema.org/Article':
        metadata.append(parse_article_meta(entry))

    return metadata

def getLocationCode(location):
    switcher = {
    	'Vancouver': 'CITY01',
    	'Richmond': 'CITY02',
    	'Burnaby': 'CITY03'
    }
    return switcher.get(location, "nothing")

def get_paged_meta(url, location, num_pages):
    collected_meta = []
    location_code = getLocationCode(location)

    for page in range(1, num_pages):
    	next_url = url + '?page=' + str(page) + '&location=' + location_code
    	CITY01_posts = get_metadata_from(next_url)
    	collected_meta += CITY01_posts

    return collected_meta


def get_article_content_meta(url, link, title, img):
    collected_meta = ''
    content_meta = ''
    content_replace = ''
    note = []
    content_url = url + link
    print(content_url)
    resp = fetch(content_url)
    post_entries = parse_article_content_entries(resp.text)

    for entry in post_entries:
    	meta = parse_article_content_meta(entry)

    	for replace_p in meta['content_replace_p']:
	    	if '©' in replace_p.text:
	    		content_replace = replace_p.text

    #把文章奇怪的字串濾掉
    content_meta = meta['content'].replace(content_replace, '').replace('联系我时请说明是在Vansky上看到的，谢谢！', '')

   	#比較關鍵字時把標題拉進去一起比，但只回傳內容
    collected_meta += title + '\n'
    collected_meta += content_meta

    if '/images/wutupian.png' in img:
    	note.append('無照片')
    else:
    	note.append('有照片')

    if '$' in collected_meta:
    	note.append('有價格')
    else:
    	if re.search('(([租|金|包|月|币])\d{3})|[\d]{2}(\/月|\/mth|\/month|\/day|包|元|刀|加|CAD)|\$[\d]', collected_meta.replace(' ', '')):
    		note.append('有價格')
    	else:
    		note.append('無價格')

    return note, meta['time'], content_meta


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    r = paragraph.add_run()
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True
    
    r._r.text = text
    hyperlink.append(r._r)
    paragraph._p.append(hyperlink)

    return hyperlink

def create_upload_file(location):
	today = datetime.date.today()
	yesterday = today - datetime.timedelta(days=1)

	path_or_stream = 'vansky_{}_{}.docx'.format(location, yesterday.strftime("%Y%m%d"))
	url = 'https://www.vansky.com/info/ZFBG08.html'
	resp = fetch(url)  # step-1

	post_entries = parse_article_entries(resp.text)
	metadata = get_paged_meta(url, location, num_pages=6)
	document = docx.Document()
	for meta in metadata:
		note, mata_time, article_content = get_article_content_meta('https://www.vansky.com/info/', meta['link'], meta['title'], meta['img'])
		p = document.add_paragraph('')
		r = p.add_run('({}) {}\n'.format(','.join(note), meta['title']))
		r.font.bold = True #加粗
		article_link = 'https://www.vansky.com/info/' + meta['link']

		s = p.add_run('{}\n'.format(mata_time))
		s.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
		add_hyperlink(p, article_link, article_link)
		p.add_run('\n{}\n'.format(article_content))

	document.save(path_or_stream)
	google_upload(path_or_stream)


def get_posts(metadata):
	with Pool(processes=3) as pool:
		contents = pool.map(create_upload_file, metadata)
		return contents


if __name__ == '__main__':
	start = time.time()
	#metadata = ['Vancouver', 'Richmond', 'Burnaby']
	metadata = ['Vancouver']
	get_posts(metadata)

	print('花費: %f 秒' % (time.time() - start))